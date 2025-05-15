from typing import List
from doc_parser.parser import DocParser, Section
import docx
from docx.enum.style import WD_STYLE_TYPE
from docx.text.paragraph import Paragraph
from docx.table import Table
import re
import json
import zipfile
import xml.etree.ElementTree as ET
from collections import defaultdict


def format_number(num_type: str, number: int) -> str:
    """
    Format a number according to a specific numbering style (e.g., decimal, roman, bullet).

    Args:
        num_type (str): Numbering style type (e.g., 'decimal', 'lowerRoman').
        number (int): The number to format.

    Returns:
        str: The formatted string (e.g., '1', 'i', 'A', '•').
    """
    def to_roman(n: int, upper=True):
        val = [
            1000, 900, 500, 400,
            100, 90, 50, 40,
            10, 9, 5, 4, 1
        ]
        syms = [
            "M", "CM", "D", "CD",
            "C", "XC", "L", "XL",
            "X", "IX", "V", "IV", "I"
        ]
        roman = ''
        for i in range(len(val)):
            count = int(n / val[i])
            roman += syms[i] * count
            n -= val[i] * count
        return roman if upper else roman.lower()

    if num_type == "bullet":
        return "•"
    elif num_type == "decimal":
        return str(number)
    elif num_type == "lowerLetter":
        return chr(ord('a') + number - 1)
    elif num_type == "upperLetter":
        return chr(ord('A') + number - 1)
    elif num_type == "lowerRoman":
        return to_roman(number, upper=False)
    elif num_type == "upperRoman":
        return to_roman(number, upper=True)
    else:
        raise ValueError("Unknown type")


def parse_numbering_xml(docx_path):
    """
    Parse the numbering.xml inside a .docx file to get the mapping of numbering styles.

    Args:
        docx_path (str): Path to the .docx file.

    Returns:
        dict: A nested dictionary of {numId: {ilvl: numFormat}}.
    """
    numbering_info = defaultdict(dict)
    with zipfile.ZipFile(docx_path, 'r') as docx:
        if 'word/numbering.xml' not in docx.namelist():
            raise Exception("This document has no numbering.xml.")

        xml_data = docx.read('word/numbering.xml')
        root = ET.fromstring(xml_data)

        ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

        abstract_nums = {}
        for abstract in root.findall('w:abstractNum', ns):
            abstract_id = abstract.attrib.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}abstractNumId')
            for lvl in abstract.findall('w:lvl', ns):
                ilvl = lvl.attrib.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ilvl')
                numFmt = lvl.find('w:numFmt', ns)
                if numFmt is not None:
                    abstract_nums[(abstract_id, ilvl)] = numFmt.attrib.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')

        for num in root.findall('w:num', ns):
            num_id = num.attrib.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numId')
            abstract_num_id_elem = num.find('w:abstractNumId', ns)
            if abstract_num_id_elem is not None:
                abstract_id = abstract_num_id_elem.attrib.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
                for ilvl in range(0, 10):
                    key = (abstract_id, str(ilvl))
                    if key in abstract_nums:
                        numbering_info[num_id][ilvl] = abstract_nums[key]

    return numbering_info


def add_paragraph(index, para, heading_styles, hierarchy, numbering):
    """
    Analyze and append a paragraph's metadata to the hierarchy list.

    Args:
        index (int): Paragraph index in the document.
        para (Paragraph): The docx paragraph object.
        heading_styles (dict): Mapping of heading style names to levels.
        hierarchy (list): List to collect paragraph structure info.
        numbering (str): Optional numbering prefix.
    """
    level = 0
    is_heading = False
    indent = 0
    
    if para.style.name in heading_styles:
        level = heading_styles[para.style.name]
        is_heading = True
    
    if hasattr(para.paragraph_format, 'left_indent') and para.paragraph_format.left_indent:
        indent = para.paragraph_format.left_indent.pt
        indent_level = int(indent / 36)
        if not is_heading and indent_level > 0:
            level = indent_level

    hierarchy.append({
        'index': index,
        'text': f"{numbering}. {para.text}" if numbering else para.text,
        'level': level,
        'is_heading': is_heading,
        'style': para.style.name,
        'indent': indent
    })


def get_content(table):
    """
    Convert a docx table into a flat string representation.

    Args:
        table (Table): A docx table object.

    Returns:
        str: Flattened textual representation of the table.
    """
    table_text = []
    for row in table.rows:
        row_text = []
        for cell in row.cells:
            row_text.append(cell.text.strip())
        table_text.append(", ".join(row_text))
        table_text.append("\n")

    return " ".join(table_text)


def add_table(index, table, heading_styles, hierarchy):
    """
    Add table content to the hierarchy as a pseudo-paragraph node.

    Args:
        index (int): Index in the document.
        table (Table): The docx table object.
        heading_styles (dict): Heading styles for context.
        hierarchy (list): List to collect hierarchy nodes.
    """
    level = 0 if not hierarchy else hierarchy[-1]['level']
    indent = 0 if not hierarchy else hierarchy[-1]['indent']
    hierarchy.append({
        'index': index,
        'text': get_content(table),
        'level': level,
        'is_heading': False,
        'style': 'table',
        'indent': indent
    })


def extract_paragraph_hierarchy(docx_path):
    """
    Extracts all paragraphs and tables from a docx file and builds a flat hierarchy
    based on indentation and heading levels.

    Args:
        docx_path (str): Path to the .docx file.

    Returns:
        list: A list of dicts representing paragraphs or tables with structure info.
    """
    numbering_style = parse_numbering_xml(docx_path)
    numbering_style_counter = {}
    doc = docx.Document(docx_path)

    hierarchy = []

    heading_styles = {}
    for style in doc.styles:
        if style.name.startswith('Heading'):
            try:
                level = int(style.name.split(' ')[-1])
                heading_styles[style.name] = level
            except:
                pass
    

    for index, item in enumerate(doc.iter_inner_content()):
        if isinstance(item, Paragraph):
            para = item
            if not para.text.strip():
                continue
            if para._element.pPr is not None and para._element.pPr.numPr is not None:
                num_id = str(para._element.pPr.numPr.numId.val)
                ilvl = para._element.pPr.numPr.ilvl.val
                style_type = numbering_style[num_id][ilvl]
                counter_key = f"{num_id}&{ilvl}"
                numbering_style_counter[counter_key] = numbering_style_counter.get(counter_key, 0) + 1
                numbering = format_number(style_type, numbering_style_counter[counter_key])
            else:
                numbering = ""
            add_paragraph(index, para, heading_styles, hierarchy, numbering)
        elif isinstance(item, Table):
            table = item
            add_table(index, table, heading_styles, hierarchy)
    
    return hierarchy


def get_structured_content(hierarchy):
    """
    Convert a flat list of paragraphs/tables with levels into a nested tree structure.

    Args:
        hierarchy (list): Flat list of elements with 'level' and 'text'.

    Returns:
        list: Nested structure with 'children' fields indicating sublevels.
    """
    result = []
    stack = [(0, result)]  # (level, children_list)

    for item in hierarchy:
        current_level = item['level']
        node = {'level': current_level, 'text': item['text'], 'children': []}

        while stack and stack[-1][0] >= current_level:
            stack.pop()
            
        if not stack:
            stack.append((0, result))
            
        stack[-1][1].append(node)
        stack.append((current_level, node['children']))
    
    return result


def format_content(data):
    """
    Recursively format the nested structure into a flat string.

    Args:
        data (list): Structured content with children.

    Returns:
        str: Concatenated text content from all nested nodes.
    """
    content = ''
    for item in data:
        content += item['text'] + " " + format_content(item['children'])
    return content


class DocxParser(DocParser):
    def section_parse(self) -> List[Section]:
        """
        Extract structured sections from a .docx file by analyzing paragraph hierarchy.

        Returns:
            List[Section]: A list of Section objects, each with a title and corresponding content.
        """
        hierarchy = extract_paragraph_hierarchy(self.file_path)
        structured_content = get_structured_content(hierarchy)
        final = []
        for item in structured_content:
            if item['level'] != 0 or item['children']:
                final.append(item)

        result: List[Section] = []
        for item in final:
            title = item['text']
            content = format_content(item['children'])
            result.append(Section(title=title, content=content))
        return result
